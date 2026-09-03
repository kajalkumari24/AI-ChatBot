import io
import csv
import zipfile
from PIL import Image
from pypdf import PdfReader, PdfWriter
from pdf2image import convert_from_bytes
from docx import Document
import openpyxl
import pandas as pd
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter


# ============================================================
# IMAGE CONVERSIONS
# ============================================================

def convert_image_format(image_bytes: bytes, target_format: str) -> bytes:
    """target_format: 'JPEG', 'PNG', 'WEBP'"""
    img = Image.open(io.BytesIO(image_bytes))
    if target_format.upper() == "JPEG" and img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format=target_format.upper())
    return buf.getvalue()


def resize_image(image_bytes: bytes, width: int, height: int) -> bytes:
    img = Image.open(io.BytesIO(image_bytes))
    img = img.resize((width, height))
    buf = io.BytesIO()
    img.save(buf, format=img.format or "PNG")
    return buf.getvalue()


def compress_image(image_bytes: bytes, quality: int = 50) -> bytes:
    img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=quality, optimize=True)
    return buf.getvalue()


def rotate_image(image_bytes: bytes, degrees: int) -> bytes:
    img = Image.open(io.BytesIO(image_bytes))
    rotated = img.rotate(-degrees, expand=True)
    buf = io.BytesIO()
    rotated.save(buf, format=img.format or "PNG")
    return buf.getvalue()


def crop_image(image_bytes: bytes, left: int, top: int, right: int, bottom: int) -> bytes:
    img = Image.open(io.BytesIO(image_bytes))
    cropped = img.crop((left, top, right, bottom))
    buf = io.BytesIO()
    cropped.save(buf, format=img.format or "PNG")
    return buf.getvalue()


# ============================================================
# IMAGE <-> PDF
# ============================================================

def images_to_pdf(images_bytes: list[bytes]) -> bytes:
    """Combines one or more images into a single PDF, one image per page."""
    pil_images = []
    for img_bytes in images_bytes:
        img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
        pil_images.append(img)

    buf = io.BytesIO()
    if len(pil_images) == 1:
        pil_images[0].save(buf, format="PDF")
    else:
        pil_images[0].save(buf, format="PDF", save_all=True, append_images=pil_images[1:])
    return buf.getvalue()


def pdf_to_images(pdf_bytes: bytes, image_format: str = "PNG") -> list[tuple[str, bytes]]:
    """Returns a list of (filename, image_bytes) — one per PDF page."""
    pages = convert_from_bytes(pdf_bytes)
    results = []
    for i, page in enumerate(pages, start=1):
        buf = io.BytesIO()
        page.save(buf, format=image_format.upper())
        results.append((f"page_{i}.{image_format.lower()}", buf.getvalue()))
    return results


# ============================================================
# PDF TEXT & DOCUMENT CONVERSIONS
# ============================================================

def pdf_to_text(pdf_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def text_to_pdf(text: str) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    lines = text.split("\n")
    y = height - 40
    for line in lines:
        if y < 40:
            c.showPage()
            y = height - 40
        c.drawString(40, y, line[:110])  # basic line-length guard
        y -= 14
    c.save()
    return buf.getvalue()


def text_to_docx(text: str) -> bytes:
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def docx_to_text(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


def docx_to_pdf_via_text(docx_bytes: bytes) -> bytes:
    """
    Simple text-based DOCX->PDF (preserves text content, not original layout/styling).
    Full-fidelity conversion normally requires LibreOffice — out of scope for now.
    """
    text = docx_to_text(docx_bytes)
    return text_to_pdf(text)


# ============================================================
# SPREADSHEET CONVERSIONS
# ============================================================

def csv_to_xlsx(csv_bytes: bytes) -> bytes:
    text = csv_bytes.decode("utf-8")
    reader = csv.reader(io.StringIO(text))

    wb = openpyxl.Workbook()
    ws = wb.active
    for row in reader:
        ws.append(row)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def xlsx_to_csv(xlsx_bytes: bytes) -> bytes:
    df = pd.read_excel(io.BytesIO(xlsx_bytes))
    buf = io.StringIO()
    df.to_csv(buf, index=False)
    return buf.getvalue().encode("utf-8")


# ============================================================
# PDF MANIPULATION
# ============================================================

def merge_pdfs(pdf_bytes_list: list[bytes]) -> bytes:
    writer = PdfWriter()
    for pdf_bytes in pdf_bytes_list:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        for page in reader.pages:
            writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def split_pdf(pdf_bytes: bytes) -> list[tuple[str, bytes]]:
    """Splits a PDF into one file per page."""
    reader = PdfReader(io.BytesIO(pdf_bytes))
    results = []
    for i, page in enumerate(reader.pages, start=1):
        writer = PdfWriter()
        writer.add_page(page)
        buf = io.BytesIO()
        writer.write(buf)
        results.append((f"page_{i}.pdf", buf.getvalue()))
    return results


def rotate_pdf(pdf_bytes: bytes, degrees: int) -> bytes:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    for page in reader.pages:
        page.rotate(degrees)
        writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def zip_files(files: list[tuple[str, bytes]]) -> bytes:
    """Bundles multiple output files into a single downloadable zip."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        for filename, data in files:
            zf.writestr(filename, data)
    return buf.getvalue()
